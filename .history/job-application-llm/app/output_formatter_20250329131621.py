import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import datetime
import platform

class OutputFormatter:
    def __init__(self):
        # Fix the output directory path
        self.base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.output_dir = os.path.join(self.base_dir, 'output')
        
        # Ensure output directory exists
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            
        # Check if running in Docker
        self.in_docker = os.path.exists('/.dockerenv')
    
    def format_text(self, content, content_type):
        """Return formatted plain text."""
        if content_type == 'linkedin_message':
            # Ensure LinkedIn message is under 200 characters
            if len(content) > 200:
                content = content[:197] + '...'
        
        elif content_type in ['connection_email', 'hiring_manager_email']:
            # Ensure emails are reasonable length (roughly 200 words)
            words = content.split()
            if len(words) > 220:  # Give a little buffer
                content = ' '.join(words[:200]) + '...'
        
        elif content_type == 'cover_letter':
            # Ensure cover letter is reasonable length (roughly 350 words)
            words = content.split()
            if len(words) > 370:  # Give a little buffer
                content = ' '.join(words[:350]) + '...'
        
        return content
    
    def create_docx(self, content, job_data, candidate_data, content_type):
        """Create a DOCX file with the formatted content."""
        try:
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            if content_type == 'cover_letter':
                self._format_cover_letter_docx(doc, content, job_data, candidate_data)
            else:
                # For emails, use a simpler format
                self._format_email_docx(doc, content, job_data, candidate_data, content_type)
            
            # Generate filename
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            company_name = job_data['company_name'].replace(' ', '_')
            filename = f"{content_type}_{company_name}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # Save the document
            doc.save(filepath)
            
            # Return relative path for the API response
            return os.path.basename(filepath)
            
        except Exception as e:
            print(f"Error creating document: {str(e)}")
            return None
    
    def convert_to_pdf(self, docx_path):
        """Convert DOCX to PDF."""
        try:
            # Generate PDF path
            pdf_path = docx_path.replace('.docx', '.pdf')
            
            if self.in_docker:
                # For Docker environment, use alternative conversion method
                try:
                    from subprocess import run, PIPE
                    # Use libreoffice for conversion in Docker
                    cmd = ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                           os.path.dirname(pdf_path), docx_path]
                    process = run(cmd, stdout=PIPE, stderr=PIPE)
                    if process.returncode != 0:
                        print(f"LibreOffice conversion failed: {process.stderr.decode()}")
                        return None
                except Exception as e:
                    print(f"LibreOffice conversion failed: {str(e)}")
                    return None
            else:
                # For local environment, use docx2pdf
                convert(docx_path, pdf_path)
            
            # Check if conversion was successful
            if os.path.exists(pdf_path):
                return pdf_path
            return None
            
        except Exception as e:
            print(f"Error converting to PDF: {str(e)}")
            return None
    
    def _format_cover_letter_docx(self, doc, content, job_data, candidate_data):
        """Format a cover letter in DOCX."""
        try:
            # Add candidate info at the top
            candidate_info = doc.add_paragraph()
            candidate_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
            candidate_info.add_run(candidate_data['personal_info']['name']).bold = True
            candidate_info.add_run('\n' + candidate_data['personal_info']['email'])
            candidate_info.add_run('\n' + candidate_data['personal_info']['phone'])
            if candidate_data['personal_info']['linkedin']:
                candidate_info.add_run('\n' + candidate_data['personal_info']['linkedin'])
            
            # Add date
            date = doc.add_paragraph()
            date.alignment = WD_ALIGN_PARAGRAPH.LEFT
            date.add_run(datetime.datetime.now().strftime("%B %d, %Y"))
            
            
            # Add space before content
            doc.add_paragraph()
            
            # Split content into paragraphs and add them
            paragraphs = content.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.add_run(paragraph.strip())
            
            # Add signature space
            doc.add_paragraph()  # Add space
            signature = doc.add_paragraph()
            signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        except Exception as e:
            print(f"Error formatting cover letter: {str(e)}")
            raise
    
    def _format_email_docx(self, doc, content, job_data, candidate_data, content_type):
        """Format an email in DOCX."""
        # Add subject line for emails
        if content_type == 'connection_email':
            subject = f"Connection Request - {candidate_data['personal_info']['name']}"
        elif content_type == 'hiring_manager_email':
            subject = f"Application for {job_data['job_title']} Position - {candidate_data['personal_info']['name']}"
        
        subject_line = doc.add_paragraph()
        subject_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subject_line.add_run('Subject: ').bold = True
        subject_line.add_run(subject)
        
        # Add content
        doc.add_paragraph()  # Add space
        for paragraph in content.split('\n'):
            if paragraph.strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(paragraph) 