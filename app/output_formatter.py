import os
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import datetime
import tempfile
import time
from typing import Optional

from app.artifact_models import CoverLetterArtifact, EmailArtifact
from app.fast_pdf_generator import FastPDFGenerator

class OutputFormatter:
    def __init__(self):
        # Use a shared temp directory so downloads work across worker processes
        base_temp_dir = os.path.join(tempfile.gettempdir(), 'applify_output')
        self.output_dir = os.path.join(base_temp_dir, 'files')
        os.makedirs(self.output_dir, exist_ok=True)
            
        # Check if running in Docker or Heroku
        self.in_docker = os.path.exists('/.dockerenv')
        self.in_heroku = bool(os.environ.get('DYNO'))
        
        # Initialize fast PDF generator
        self.pdf_generator = FastPDFGenerator(output_dir=self.output_dir)
        
        # Clean up old files periodically
        self._cleanup_old_files()
    
    def _cleanup_old_files(self):
        """Clean up files older than 1 hour."""
        try:
            current_time = datetime.datetime.now()
            for filename in os.listdir(self.output_dir):
                file_path = os.path.join(self.output_dir, filename)
                if os.path.isfile(file_path):
                    file_time = datetime.datetime.fromtimestamp(os.path.getmtime(file_path))
                    if (current_time - file_time).total_seconds() > 900:  # 15 minutes
                        os.remove(file_path)
        except Exception as e:
            print(f"Error cleaning up old files: {str(e)}")
    
    def format_text(self, content, content_type):
        """Return formatted plain text."""
        if content_type == 'linkedin_message':
            # Ensure LinkedIn message is under 200 characters
            if len(content) > 200:
                content = content[:197] + '...'
        
        elif content_type in ['connection_email', 'hiring_manager_email']:
            # Remove subject line if present (LLM includes it in output)
            lines = content.split('\n')
            if lines and (lines[0].startswith('Subject:') or lines[0].startswith('Re:')):
                # Remove the subject line from display
                content = '\n'.join(lines[1:]).strip()

            # Ensure emails are reasonable length (roughly 200 words)
            words = content.split()
            if len(words) > 220:  # Give a little buffer
                content = ' '.join(words[:200]) + '...'
        
        elif content_type == 'cover_letter':
            # Ensure cover letter is reasonable length (roughly 350 words)
            words = content.split()
            if len(words) > 370:  # Give a little buffer
                content = ' '.join(words[:350]) + '...'
        
        return content

    def render_artifact_bundle(self, artifact, candidate_data, user_id: str, artifact_id: int):
        """Render the same validated artifact to DOCX and PDF."""
        content_type = artifact.metadata.content_type
        docx_info = self.create_docx_from_artifact(
            artifact=artifact,
            candidate_data=candidate_data,
            user_id=user_id,
            artifact_id=artifact_id,
        )
        pdf_info = self.create_pdf_from_artifact(
            artifact=artifact,
            candidate_data=candidate_data,
            user_id=user_id,
            artifact_id=artifact_id,
        )
        if not docx_info and not pdf_info:
            return None

        return {
            "artifact_id": artifact_id,
            "content_type": content_type,
            "available_formats": [
                fmt for fmt, value in (("docx", docx_info), ("pdf", pdf_info)) if value
            ],
            "docx": self._public_file_info(docx_info, "docx") if docx_info else None,
            "pdf": self._public_file_info(pdf_info, "pdf") if pdf_info else None,
        }

    def create_docx_from_artifact(self, artifact, candidate_data, user_id: str, artifact_id: int):
        """Create a DOCX file from a validated artifact."""
        try:
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            if isinstance(artifact, CoverLetterArtifact):
                self._format_cover_letter_artifact_docx(doc, artifact)
            elif isinstance(artifact, EmailArtifact):
                self._format_email_artifact_docx(doc, artifact)
            else:
                raise ValueError(f"Unsupported artifact type: {type(artifact).__name__}")

            filepath = self._artifact_path(user_id, artifact_id, "docx")
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            doc.save(filepath)
            return self._verify_file(filepath, artifact.metadata.content_type, "docx")
        except Exception as e:
            logging.error(f"Error creating structured DOCX: {str(e)}", exc_info=True)
            return None

    def create_pdf_from_artifact(self, artifact, candidate_data, user_id: str, artifact_id: int):
        """Create a PDF file from the same validated artifact."""
        try:
            filepath = self._artifact_path(user_id, artifact_id, "pdf")
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            if isinstance(artifact, CoverLetterArtifact):
                pdf_result = self.pdf_generator.generate_cover_letter_pdf(
                    content=artifact.to_plain_text(),
                    job_data=artifact.metadata.model_dump(),
                    candidate_data=candidate_data,
                    output_path=filepath,
                )
            elif isinstance(artifact, EmailArtifact):
                pdf_result = self.pdf_generator.generate_cover_letter_pdf(
                    content=artifact.to_plain_text(),
                    job_data=artifact.metadata.model_dump(),
                    candidate_data=candidate_data,
                    output_path=filepath,
                )
            else:
                raise ValueError(f"Unsupported artifact type: {type(artifact).__name__}")
            if not pdf_result:
                return None
            return self._verify_file(filepath, artifact.metadata.content_type, "pdf")
        except Exception as e:
            logging.error(f"Error creating structured PDF: {str(e)}", exc_info=True)
            return None

    def get_artifact_download_path(self, user_id: str, artifact_id: int, fmt: str) -> Optional[str]:
        """Resolve a user-scoped artifact file path."""
        if fmt not in {"docx", "pdf"}:
            return None
        filepath = self._artifact_path(user_id, artifact_id, fmt)
        if os.path.exists(filepath):
            return filepath
        return None

    def _artifact_path(self, user_id: str, artifact_id: int, fmt: str) -> str:
        safe_user_id = str(user_id).replace("/", "_")
        return os.path.join(self.output_dir, safe_user_id, str(artifact_id), f"artifact.{fmt}")

    def _verify_file(self, filepath: str, content_type: str, fmt: str):
        max_retries = 3
        for _ in range(max_retries):
            if os.path.exists(filepath) and os.path.getsize(filepath) > 0:
                return {
                    "filename": os.path.basename(filepath),
                    "filepath": filepath,
                    "size_bytes": os.path.getsize(filepath),
                    "content_type": content_type,
                    "format": fmt,
                }
            time.sleep(0.1)
        return None

    def _public_file_info(self, file_info, fmt: str):
        return {
            "filename": file_info["filename"],
            "size_bytes": file_info.get("size_bytes", 0),
            "format": fmt,
        }
    
    def create_docx(self, content, job_data, candidate_data, content_type):
        """Create a DOCX file with the formatted content."""
        try:
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            if content_type == 'cover_letter':
                self._format_cover_letter_docx(doc, content, job_data, candidate_data)
            else:
                # For emails, use a simpler format
                self._format_email_docx(doc, content, job_data, candidate_data, content_type)
            
            # Generate filename with user ID to prevent conflicts
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            # Sanitize company name - remove invalid filename characters
            company_name = (job_data.get('company_name') or 'Company').replace(' ', '_').replace('/', '_').replace('\\', '_')
            filename = f"{content_type}_{company_name}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Save the document
            doc.save(filepath)
            
            # Ensure file is written to disk and verify it exists
            max_retries = 3
            for i in range(max_retries):
                if os.path.exists(filepath) and os.path.getsize(filepath) > 0:
                    logging.info(f"Document verified on disk: {filename} ({os.path.getsize(filepath)} bytes)")
                    break
                time.sleep(0.1)  # Wait 100ms between checks
            else:
                logging.error(f"Document not found after saving: {filepath}")
                return None
            
            # Return both the filename and the full path
            return {
                'filename': filename,
                'filepath': filepath
            }
            
        except Exception as e:
            logging.error(f"Error creating document: {str(e)}", exc_info=True)
            return None
    
    def convert_to_pdf(self, docx_info):
        """
        Deprecated compatibility wrapper.
        """
        raise ValueError("Direct DOCX conversion is no longer supported; render PDFs from structured artifacts")
    
    def create_pdf_direct(self, content, job_data, candidate_data, content_type):
        """
        Create PDF directly without DOCX intermediate step.
        
        ULTRA-FAST: ~30-100ms generation time.
        """
        try:
            if content_type == 'cover_letter':
                return self.pdf_generator.generate_cover_letter_pdf(
                    content, job_data, candidate_data
                )
            elif content_type in ['connection_email', 'hiring_manager_email', 'linkedin_message']:
                # For emails, use cover letter format but with email-specific styling
                return self.pdf_generator.generate_cover_letter_pdf(
                    content, job_data, candidate_data
                )
            else:
                print(f"PDF generation not supported for content type: {content_type}")
                return None
                
        except Exception as e:
            print(f"Error creating PDF directly: {str(e)}")
            return None
    
    def create_resume_pdf_direct(self, resume_data, candidate_data, job_title="Position"):
        """
        Create optimized resume PDF directly - ULTRA-FAST.
        
        Performance: ~50-100ms for complete resume PDF.
        """
        try:
            return self.pdf_generator.generate_resume_pdf(
                resume_data, candidate_data, job_title
            )
        except Exception as e:
            print(f"Error creating resume PDF: {str(e)}")
            return None
    
    def _format_cover_letter_docx(self, doc, content, job_data, candidate_data):
        """Format a cover letter in DOCX."""
        try:
            personal_info = candidate_data.get('personal_info', {})
            header_parts = []
            name = personal_info.get('name')
            email = personal_info.get('email')
            phone = personal_info.get('phone')
            location = personal_info.get('location')

            placeholder_info = {
                'name': 'Candidate Name',
                'email': 'candidate@email.com',
                'phone': '(555) 123-4567',
                'location': 'City, State'
            }

            def is_placeholder(value, key):
                return value is None or value == '' or value == placeholder_info[key]

            has_real_info = any(
                not is_placeholder(personal_info.get(key), key)
                for key in placeholder_info
            )

            if has_real_info:
                # Add name line (bold)
                if not is_placeholder(name, 'name'):
                    name_para = doc.add_paragraph()
                    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    name_run = name_para.add_run(name)
                    name_run.bold = True

                if not is_placeholder(email, 'email'):
                    header_parts.append(email)
                if not is_placeholder(phone, 'phone'):
                    header_parts.append(phone)
                if not is_placeholder(location, 'location'):
                    header_parts.append(location)

                if header_parts:
                    contact_para = doc.add_paragraph()
                    contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    contact_para.add_run(' - '.join(header_parts))

                # Add space before LLM content
                doc.add_paragraph()

            # Split content into paragraphs and add them
            paragraphs = content.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Set font to Times New Roman 12pt
                    run = p.add_run(paragraph.strip())
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

                    # Set 1.5 line spacing
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            
            # # Add signature space
            # doc.add_paragraph()  # Add space
            # signature = doc.add_paragraph()
            # signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        except Exception as e:
            print(f"Error formatting cover letter: {str(e)}")
            raise
    
    def _format_email_docx(self, doc, content, job_data, candidate_data, content_type):
        """Format an email in DOCX."""
        # LLM already generates subject line, so just add the content
        # Split content into paragraphs and add them
        if content:
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.add_run(paragraph) 

    def _format_cover_letter_artifact_docx(self, doc, artifact: CoverLetterArtifact):
        self._add_header_lines(doc, artifact.header)
        self._add_paragraph_line(doc, artifact.greeting, font_name="Times New Roman")
        self._add_paragraph_line(doc, artifact.opening.text, font_name="Times New Roman")
        for block in artifact.body:
            self._add_paragraph_line(doc, block.text, font_name="Times New Roman")
        self._add_paragraph_line(doc, artifact.closing.text, font_name="Times New Roman")
        self._add_paragraph_line(doc, artifact.signature.signoff, font_name="Times New Roman")

    def _format_email_artifact_docx(self, doc, artifact: EmailArtifact):
        self._add_header_lines(doc, artifact.header, include_name=False)
        self._add_paragraph_line(doc, f"Subject: {artifact.subject}")
        self._add_paragraph_line(doc, artifact.greeting)
        for block in artifact.body:
            self._add_paragraph_line(doc, block.text)
        self._add_paragraph_line(doc, artifact.call_to_action.text)
        self._add_paragraph_line(doc, artifact.signature.signoff)

    def _add_header_lines(self, doc, header, include_name=True):
        header_parts = [
            part for part in [
                header.applicant_email,
                header.applicant_phone,
                header.applicant_location,
            ] if part
        ]
        if include_name and header.applicant_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(header.applicant_name)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        if header_parts:
            self._add_paragraph_line(doc, " - ".join(header_parts), font_size=11, spacing=False)
        if include_name or header_parts:
            doc.add_paragraph()

    def _add_paragraph_line(self, doc, text, font_name="Calibri", font_size=12, spacing=True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if spacing:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
