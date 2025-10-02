import os
import json
import re
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn
import datetime
import tempfile

class ResumeRefiner:
    """Advanced resume refinement system that tailors resumes to job descriptions."""
    
    def __init__(self):
        self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        # ATS-friendly formatting standards
        self.formatting_standards = {
            'margins': 1.0,
            'font_name': 'Times New Roman',
            'font_size_body': 11,
            'font_size_name': 14,
            'font_size_headers': 12,
            'line_spacing': 1.15,
            'max_page_length': 1,
            'max_word_count': 600,
            'section_order': [
                'contact_info',
                'professional_summary', 
                'technical_skills',
                'professional_experience',
                'education',
                'certifications'
            ]
        }
        
        # Modern resume templates
        self.resume_templates = {
            'technical': {
                'emphasis': ['technical_skills', 'professional_experience', 'education'],
                'summary_style': 'technical_focused',
                'skills_format': 'categorized'
            },
            'business': {
                'emphasis': ['professional_experience', 'achievements', 'education'],
                'summary_style': 'results_focused',
                'skills_format': 'comprehensive'
            },
            'creative': {
                'emphasis': ['portfolio', 'professional_experience', 'creative_skills'],
                'summary_style': 'creative_focused',
                'skills_format': 'integrated'
            }
        }
    
    def _parse_json_response(self, response_text):
        """Safely parse JSON response from LLM, handling markdown and errors."""
        try:
            # Clean markdown code blocks
            cleaned = response_text.strip()
            cleaned = re.sub(r'```json\s*', '', cleaned)
            cleaned = re.sub(r'```\s*$', '', cleaned)
            cleaned = cleaned.strip()
            
            # Check for HTML error responses
            if cleaned.startswith('<') or cleaned.startswith('<!'):
                raise ValueError(f"Received HTML response instead of JSON: {cleaned[:100]}")
            
            # Validate JSON structure
            if not (cleaned.startswith('{') or cleaned.startswith('[')):
                raise ValueError(f"Response doesn't look like JSON: {cleaned[:100]}")
            
            return json.loads(cleaned)
        except json.JSONDecodeError as e:
            raise ValueError(f"Failed to parse JSON: {str(e)}. Response: {response_text[:200]}")
    
    def analyze_job_requirements(self, job_description):
        """Analyze job description to extract key requirements and keywords."""
        prompt = f"""
        Analyze this job description and extract key information for resume optimization:

        Job Description:
        {job_description}

        Return a JSON object with:
        {{
            "job_title": "extracted job title",
            "industry": "industry/field",
            "required_skills": ["skill1", "skill2", ...],
            "preferred_skills": ["skill1", "skill2", ...],
            "key_qualifications": ["qualification1", "qualification2", ...],
            "company_values": ["value1", "value2", ...],
            "job_level": "entry/mid/senior",
            "resume_template": "technical/business/creative",
            "important_keywords": ["keyword1", "keyword2", ...],
            "years_experience": "number or range",
            "education_requirements": "requirements",
            "certifications": ["cert1", "cert2", ...],
            "hard_skills": ["skill1", "skill2", ...],
            "achievement_focus": ["revenue", "efficiency", "innovation", "leadership", ...]
        }}
        """
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert job market analyst. Extract precise information from job descriptions for resume optimization. Return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            
            response_text = response.choices[0].message.content.strip()
            return self._parse_json_response(response_text)
        except Exception as e:
            print(f"Error analyzing job requirements: {str(e)}")
            return self._get_default_job_analysis()
    
    def _get_default_job_analysis(self):
        """Return default job analysis if parsing fails."""
        return {
            "job_title": "Professional Role",
            "industry": "Technology",
            "required_skills": [],
            "preferred_skills": [],
            "key_qualifications": [],
            "company_values": [],
            "job_level": "mid",
            "resume_template": "business",
            "important_keywords": [],
            "years_experience": "2-5",
            "education_requirements": "Bachelor's degree",
            "certifications": [],
            "hard_skills": [],
            "achievement_focus": ["results", "efficiency"]
        }
    
    def analyze_current_resume(self, candidate_data):
        """Analyze current resume content for optimization opportunities."""
        prompt = f"""
        Analyze this resume data and provide optimization insights:

        Resume Data:
        {json.dumps(candidate_data, indent=2)}

        Return a JSON object with:
        {{
            "current_strengths": ["strength1", "strength2", ...],
            "improvement_areas": ["area1", "area2", ...],
            "keyword_gaps": ["missing_keyword1", "missing_keyword2", ...],
            "experience_relevance": "high/medium/low",
            "skills_match": "high/medium/low",
            "achievements_quantified": true/false,
            "ats_optimization_score": 85,
            "formatting_issues": ["issue1", "issue2", ...],
            "content_suggestions": ["suggestion1", "suggestion2", ...],
            "recommended_reorder": ["section1", "section2", ...],
            "word_count_status": "optimal/too_long/too_short"
        }}
        """
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert resume analyst. Analyze resume content for optimization opportunities. Return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            
            response_text = response.choices[0].message.content.strip()
            return self._parse_json_response(response_text)
        except Exception as e:
            print(f"Error analyzing resume: {str(e)}")
            return {"current_strengths": [], "improvement_areas": [], "ats_optimization_score": 70}
    
    def quick_resume_analysis(self, candidate_data, job_analysis):
        """Fast resume analysis using local logic instead of API calls."""
        try:
            resume_data = candidate_data.get('resume', {})
            candidate_skills = set(skill.lower() for skill in resume_data.get('skills', []))
            required_skills = set(skill.lower() for skill in job_analysis.get('required_skills', []))
            
            # Calculate skills match
            skills_overlap = candidate_skills & required_skills
            skills_match = "high" if len(skills_overlap) >= 3 else "medium" if len(skills_overlap) >= 1 else "low"
            
            # Calculate base optimization score
            base_score = 70
            if skills_match == "high":
                base_score += 20
            elif skills_match == "medium":
                base_score += 10
            
            # Check for quantified achievements
            achievements_quantified = any(
                any(char.isdigit() for char in exp.get('description', ''))
                for exp in resume_data.get('experience', [])
            )
            if achievements_quantified:
                base_score += 10
            
            return {
                "current_strengths": list(skills_overlap)[:3],
                "improvement_areas": ["keyword optimization", "quantified achievements"],
                "keyword_gaps": list(required_skills - candidate_skills)[:5],
                "skills_match": skills_match,
                "experience_relevance": "medium",
                "achievements_quantified": achievements_quantified,
                "ats_optimization_score": min(base_score, 100),
                "formatting_issues": [],
                "content_suggestions": ["Add more quantified achievements", "Include relevant keywords"],
                "recommended_reorder": ["professional_summary", "skills", "experience", "education"],
                "word_count_status": "optimal"
            }
        except Exception as e:
            print(f"Error in quick resume analysis: {str(e)}")
            return {"current_strengths": [], "improvement_areas": [], "ats_optimization_score": 70}
    
    def generate_optimized_resume(self, candidate_data, job_analysis, resume_analysis):
        """Generate an optimized resume tailored to the job description."""
        
        # Select appropriate template
        template = self.resume_templates.get(job_analysis['resume_template'], self.resume_templates['business'])
        
        # Generate optimized content for each section
        optimized_sections = {}
        
        # Professional Summary
        optimized_sections['professional_summary'] = self._generate_professional_summary(
            candidate_data, job_analysis, resume_analysis
        )
        
        # Technical/Professional Skills
        optimized_sections['skills'] = self._optimize_skills_section(
            candidate_data, job_analysis, template
        )
        
        # Professional Experience
        optimized_sections['experience'] = self._optimize_experience_section(
            candidate_data, job_analysis, resume_analysis
        )
        
        # Education
        optimized_sections['education'] = self._optimize_education_section(
            candidate_data, job_analysis
        )
        
        return {
            'sections': optimized_sections,
            'template': template,
            'optimization_score': self._calculate_optimization_score(job_analysis, resume_analysis),
            'formatting_rules': self.formatting_standards,
            'word_count': self._estimate_word_count(optimized_sections)
        }
    
    def generate_optimized_resume_fast(self, candidate_data, job_analysis, resume_analysis):
        """Generate optimized resume with minimal API calls for speed."""
        try:
            # Use only one API call for the most critical section
            optimized_summary = self._generate_professional_summary(
                candidate_data, job_analysis, resume_analysis
            )
            
            # Use local optimization for other sections to avoid API delays
            optimized_skills = self._optimize_skills_local(candidate_data, job_analysis)
            optimized_experience = self._optimize_experience_local(candidate_data, job_analysis)
            optimized_education = self._optimize_education_section(candidate_data, job_analysis)
            
            optimized_sections = {
                'professional_summary': optimized_summary,
                'skills': optimized_skills,
                'experience': optimized_experience,
                'education': optimized_education
            }
            
            template = self.resume_templates.get(job_analysis.get('resume_template', 'business'), self.resume_templates['business'])
            
            return {
                'sections': optimized_sections,
                'template': template,
                'optimization_score': resume_analysis.get('ats_optimization_score', 75),
                'formatting_rules': self.formatting_standards,
                'word_count': self._estimate_word_count(optimized_sections)
            }
        except Exception as e:
            print(f"Error in fast resume generation: {str(e)}")
            return self._get_fallback_resume(candidate_data, job_analysis)
    
    def _optimize_skills_local(self, candidate_data, job_analysis):
        """Local skills optimization without API calls."""
        try:
            candidate_skills = candidate_data['resume'].get('skills', [])
            required_skills = job_analysis.get('required_skills', [])
            preferred_skills = job_analysis.get('preferred_skills', [])
            
            # Prioritize skills that match job requirements
            technical_skills = []
            
            # Add matching required skills first
            for skill in candidate_skills:
                if any(req.lower() in skill.lower() for req in required_skills):
                    technical_skills.append(skill)
                elif any(pref.lower() in skill.lower() for pref in preferred_skills):
                    technical_skills.append(skill)
                else:
                    technical_skills.append(skill)
            
            return {
                "technical_skills": technical_skills[:8],
                "tools_technologies": [skill for skill in technical_skills if any(tech in skill.lower() for tech in ['python', 'sql', 'javascript', 'aws', 'docker', 'kubernetes'])][:6],
                "certifications": []
            }
        except Exception as e:
            print(f"Error in local skills optimization: {str(e)}")
            return {
                "technical_skills": candidate_data['resume']['skills'][:8],
                "tools_technologies": [],
                "certifications": []
            }
    
    def _optimize_experience_local(self, candidate_data, job_analysis):
        """Local experience optimization without API calls."""
        try:
            experiences = candidate_data['resume'].get('experience', [])
            required_skills = [skill.lower() for skill in job_analysis.get('required_skills', [])]
            
            optimized_experiences = []
            for exp in experiences[:4]:  # Top 4 experiences
                # Enhance description with keywords
                description = exp.get('description', '')
                
                # Create bullet points from description
                sentences = [s.strip() for s in description.split('.') if s.strip()]
                bullet_points = []
                
                for sentence in sentences[:3]:  # Max 3 bullets per job
                    # Add quantification if missing
                    if not any(char.isdigit() for char in sentence):
                        sentence = f"• Successfully {sentence.lower()}, contributing to team efficiency"
                    else:
                        sentence = f"• {sentence}"
                    bullet_points.append(sentence)
                
                # Ensure at least 2 bullet points
                if len(bullet_points) < 2:
                    bullet_points.append(f"• Demonstrated expertise in {', '.join(required_skills[:2])} while maintaining high performance standards")
                
                optimized_exp = {
                    "company": exp.get('company', ''),
                    "title": exp.get('title', ''),
                    "start_date": exp.get('start_date', ''),
                    "end_date": exp.get('end_date', ''),
                    "location": exp.get('location', ''),
                    "description": description,
                    "bullet_points": bullet_points
                }
                optimized_experiences.append(optimized_exp)
            
            return optimized_experiences
        except Exception as e:
            print(f"Error in local experience optimization: {str(e)}")
            return candidate_data['resume'].get('experience', [])[:4]
    
    def quick_job_analysis(self, job_description):
        """Ultra-fast job analysis using regex and keywords - NO API calls."""
        try:
            import re
            
            # Extract job title using common patterns
            title_patterns = [
                r'(?:position|role|job)(?:\s+for)?\s+([A-Z][^.]*?)(?:\s+at|\s+\-|\n)',
                r'([A-Z][^.]*?)\s+(?:position|role|opportunity)',
                r'hiring\s+(?:a\s+)?([A-Z][^.]*?)(?:\s+to|\s+at|\n)'
            ]
            
            job_title = "Professional Role"
            for pattern in title_patterns:
                match = re.search(pattern, job_description, re.IGNORECASE)
                if match:
                    job_title = match.group(1).strip()
                    break
            
            # Extract skills using keyword matching
            skill_keywords = [
                'python', 'javascript', 'java', 'sql', 'react', 'node.js', 'aws', 'docker', 
                'kubernetes', 'machine learning', 'data analysis', 'project management',
                'agile', 'scrum', 'communication', 'leadership', 'teamwork', 'problem solving',
                'git', 'html', 'css', 'api', 'database', 'cloud', 'devops', 'ci/cd'
            ]
            
            required_skills = []
            preferred_skills = []
            important_keywords = []
            
            description_lower = job_description.lower()
            for skill in skill_keywords:
                if skill in description_lower:
                    if any(req_word in description_lower for req_word in ['required', 'must', 'essential']):
                        required_skills.append(skill.title())
                    else:
                        preferred_skills.append(skill.title())
                    important_keywords.append(skill.title())
            
            # Determine job level
            job_level = "mid"
            if any(word in description_lower for word in ['senior', 'lead', 'principal', '5+ years', '7+ years']):
                job_level = "senior"
            elif any(word in description_lower for word in ['junior', 'entry', 'graduate', '0-2 years']):
                job_level = "entry"
            
            # Determine template based on keywords
            template = "business"
            if any(word in description_lower for word in ['engineer', 'developer', 'technical', 'software']):
                template = "technical"
            elif any(word in description_lower for word in ['design', 'creative', 'marketing', 'brand']):
                template = "creative"
            
            return {
                "job_title": job_title,
                "industry": "Technology",
                "required_skills": required_skills[:8],
                "preferred_skills": preferred_skills[:6],
                "key_qualifications": [f"{job_level.title()} level experience", "Relevant technical skills"],
                "company_values": ["Innovation", "Collaboration", "Excellence"],
                "job_level": job_level,
                "resume_template": template,
                "important_keywords": important_keywords[:10],
                "years_experience": "3-5",
                "education_requirements": "Bachelor's degree preferred",
                "certifications": [],
                "hard_skills": required_skills[:5],
                "achievement_focus": ["results", "efficiency", "innovation"]
            }
        except Exception as e:
            print(f"Error in quick job analysis: {str(e)}")
            return self._get_default_job_analysis()
    
    def generate_optimized_resume_ultra_fast(self, candidate_data, job_analysis, resume_analysis):
        """Ultra-fast resume generation with NO API calls."""
        try:
            # Generate summary locally
            candidate_skills = candidate_data['resume']['skills'][:3]
            job_title = job_analysis.get('job_title', 'Professional Role')
            years_exp = job_analysis.get('years_experience', '3-5')
            
            summary = f"Results-driven professional with {years_exp} years of experience in {', '.join(candidate_skills)}. Proven track record of delivering high-quality solutions and driving organizational success. Seeking to leverage expertise in {candidate_skills[0] if candidate_skills else 'technology'} for the {job_title} role."
            
            # Optimize skills locally
            optimized_skills = self._optimize_skills_local(candidate_data, job_analysis)
            
            # Optimize experience locally
            optimized_experience = self._optimize_experience_local(candidate_data, job_analysis)
            
            # Optimize education
            optimized_education = self._optimize_education_section(candidate_data, job_analysis)
            
            optimized_sections = {
                'professional_summary': summary,
                'skills': optimized_skills,
                'experience': optimized_experience,
                'education': optimized_education
            }
            
            template = self.resume_templates.get(job_analysis.get('resume_template', 'business'), self.resume_templates['business'])
            
            # Calculate optimization score locally
            base_score = 75
            skills_boost = len(set(candidate_data['resume']['skills']) & set(job_analysis.get('required_skills', []))) * 5
            final_score = min(base_score + skills_boost, 100)
            
            return {
                'sections': optimized_sections,
                'template': template,
                'optimization_score': final_score,
                'formatting_rules': self.formatting_standards,
                'word_count': self._estimate_word_count(optimized_sections)
            }
        except Exception as e:
            print(f"Error in ultra fast resume generation: {str(e)}")
            return self._get_fallback_resume(candidate_data, job_analysis)
    
    def _get_fallback_resume(self, candidate_data, job_analysis):
        """Fallback resume structure if optimization fails."""
        return {
            'sections': {
                'professional_summary': f"Experienced professional with expertise in {', '.join(candidate_data['resume']['skills'][:3])}",
                'skills': {
                    "technical_skills": candidate_data['resume']['skills'][:8],
                    "tools_technologies": [],
                    "certifications": []
                },
                'experience': candidate_data['resume']['experience'][:4],
                'education': candidate_data['resume']['education'][:2]
            },
            'template': self.resume_templates['business'],
            'optimization_score': 70,
            'formatting_rules': self.formatting_standards,
            'word_count': 400
        }
    
    def _generate_professional_summary(self, candidate_data, job_analysis, resume_analysis):
        """Generate an optimized professional summary."""
        prompt = f"""
        Create a professional summary for this resume that perfectly matches the job requirements:

        Job Analysis: {json.dumps(job_analysis, indent=2)}
        Current Experience: {json.dumps(candidate_data['resume']['experience'][:3], indent=2)}
        Current Skills: {candidate_data['resume']['skills']}
        Current Summary: {candidate_data['resume'].get('summary', '')}

        Requirements:
        - 2-3 sentences maximum
        - Include {job_analysis['years_experience']} years experience mention
        - Incorporate 3-4 keywords from: {job_analysis['important_keywords']}
        - Highlight top 2 relevant skills: {job_analysis['required_skills'][:2]}
        - Show value proposition for {job_analysis['job_title']} role
        - Professional, confident tone
        - ATS-optimized with natural keyword placement

        Return only the professional summary text, no additional formatting.
        """
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a senior career coach specializing in ATS-optimized resume writing. Create compelling professional summaries that get results."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                timeout=30  # 30 second timeout
            )
            
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"Error generating summary: {str(e)}")
            return f"Experienced professional with expertise in {', '.join(candidate_data['resume']['skills'][:3])}."
    
    def _optimize_skills_section(self, candidate_data, job_analysis, template):
        """Optimize skills section based on job requirements."""
        prompt = f"""
        Optimize the skills section for this job application:

        Current Skills: {candidate_data['resume']['skills']}
        Required Skills: {job_analysis['required_skills']}
        Preferred Skills: {job_analysis['preferred_skills']}
        Hard Skills: {job_analysis['hard_skills']}
        Template Style: {template['skills_format']}

        Create an optimized skills section that:
        - Prioritizes required skills that candidate has
        - Includes all matching hard skills
        - Orders by relevance to job
        - Uses exact keywords from job description
        - Categorizes skills if template requires it

        Return JSON with:
        {{
            "technical_skills": ["skill1", "skill2", ...],
            "tools_technologies": ["tool1", "tool2", ...],
            "certifications": ["cert1", "cert2", ...]
        }}
        """
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert in ATS optimization and skills categorization. Create skills sections that perfectly match job requirements."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            
            response_text = response.choices[0].message.content.strip()
            return self._parse_json_response(response_text)
        except Exception as e:
            print(f"Error optimizing skills: {str(e)}")
            return {
                "technical_skills": candidate_data['resume']['skills'][:8],
                "tools_technologies": [],
                "certifications": []
            }
    
    def _optimize_experience_section(self, candidate_data, job_analysis, resume_analysis):
        """Optimize professional experience section."""
        optimized_experiences = []
        
        for exp in candidate_data['resume']['experience'][:4]:  # Top 4 most recent
            prompt = f"""
            Optimize this work experience entry for the target job:

            Current Experience: {json.dumps(exp, indent=2)}
            Target Job: {job_analysis['job_title']} in {job_analysis['industry']}
            Required Skills: {job_analysis['required_skills']}
            Important Keywords: {job_analysis['important_keywords']}
            Achievement Focus: {job_analysis['achievement_focus']}

            Requirements:
            - Keep the same company and title
            - Rewrite description with 2-4 bullet points
            - Include quantifiable achievements where possible
            - Use action verbs and power words
            - Incorporate relevant keywords naturally
            - Focus on results and impact
            - Each bullet point: 1-2 lines maximum
            - Show progression of responsibility

            Return JSON with:
            {{
                "company": "same as input",
                "title": "same as input", 
                "start_date": "same as input",
                "end_date": "same as input",
                "location": "same as input",
                "description": "Enhanced description paragraph",
                "bullet_points": ["• Achievement 1", "• Achievement 2", ...]
            }}
            """
            
            try:
                response = self.client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "You are an expert resume writer specializing in achievement-focused job descriptions. Create compelling experience entries that highlight quantifiable results."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3
                )
                
                response_text = response.choices[0].message.content.strip()
                optimized_exp = self._parse_json_response(response_text)
                optimized_experiences.append(optimized_exp)
            except Exception as e:
                print(f"Error optimizing experience: {str(e)}")
                # Keep original if optimization fails
                optimized_experiences.append(exp)
        
        return optimized_experiences
    
    def _optimize_education_section(self, candidate_data, job_analysis):
        """Optimize education section based on job requirements."""
        education_data = candidate_data['resume'].get('education', [])
        
        # Prioritize relevant education
        relevant_education = []
        for edu in education_data:
            # Add relevance scoring and reordering logic
            relevant_education.append({
                'institution': edu.get('institution', ''),
                'degree': edu.get('degree', ''),
                'field': edu.get('field', ''),
                'graduation_year': edu.get('graduation_year', ''),
                'gpa': edu.get('gpa', '') if float(edu.get('gpa', 0)) >= 3.5 else '',
                'relevant_coursework': edu.get('relevant_coursework', [])
            })
        
        return relevant_education[:2]  # Maximum 2 education entries
    
    def _calculate_optimization_score(self, job_analysis, resume_analysis):
        """Calculate overall optimization score."""
        base_score = resume_analysis.get('ats_optimization_score', 70)
        
        # Adjustments based on keyword matching
        keyword_boost = min(len(job_analysis['important_keywords']) * 2, 20)
        
        # Skills matching boost
        skills_boost = 10 if resume_analysis.get('skills_match') == 'high' else 5
        
        final_score = min(base_score + keyword_boost + skills_boost, 100)
        return final_score
    
    def _estimate_word_count(self, sections):
        """Estimate total word count of optimized resume."""
        total_words = 0
        
        for section_name, content in sections.items():
            if isinstance(content, str):
                total_words += len(content.split())
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, dict):
                        for value in item.values():
                            if isinstance(value, str):
                                total_words += len(value.split())
                            elif isinstance(value, list):
                                for v in value:
                                    if isinstance(v, str):
                                        total_words += len(v.split())
        
        return total_words
    
    def create_formatted_resume_docx(self, optimized_resume, candidate_data, job_title="", output_dir=None):
        """Create a professionally formatted DOCX resume."""
        try:
            doc = Document()
            
            # Set document margins and formatting
            self._set_document_formatting(doc)
            
            # Add sections in optimal order
            self._add_contact_header(doc, candidate_data)
            self._add_professional_summary(doc, optimized_resume['sections']['professional_summary'])
            self._add_skills_section(doc, optimized_resume['sections']['skills'])
            self._add_experience_section(doc, optimized_resume['sections']['experience'])
            self._add_education_section(doc, optimized_resume['sections']['education'])
            
            # Generate filename
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            company_name = job_title.replace(' ', '_') if job_title else 'Optimized'
            filename = f"Resume_Optimized_{company_name}_{timestamp}.docx"
            
            # Use provided output_dir or fallback to creating new OutputFormatter
            if output_dir:
                filepath = os.path.join(output_dir, filename)
            else:
                from app.output_formatter import OutputFormatter
                output_formatter = OutputFormatter()
                filepath = os.path.join(output_formatter.output_dir, filename)
            
            doc.save(filepath)
            
            return {
                'filename': filename,
                'filepath': filepath,
                'word_count': optimized_resume['word_count'],
                'optimization_score': optimized_resume['optimization_score']
            }
            
        except Exception as e:
            print(f"Error creating formatted resume: {str(e)}")
            return None
    
    def _set_document_formatting(self, doc):
        """Set standard document formatting."""
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(self.formatting_standards['margins'])
            section.bottom_margin = Inches(self.formatting_standards['margins'])
            section.left_margin = Inches(self.formatting_standards['margins'])
            section.right_margin = Inches(self.formatting_standards['margins'])
        
        # Ensure default styles use Times New Roman
        styles = doc.styles
        if 'Normal' in styles:
            normal_style = styles['Normal']
            normal_style.font.name = self.formatting_standards['font_name']
            normal_style.font.size = Pt(self.formatting_standards['font_size_body'])
        for style_name in ['List Bullet', 'List Paragraph', 'Heading 1', 'Heading 2', 'Heading 3']:
            if style_name in styles:
                style = styles[style_name]
                style.font.name = self.formatting_standards['font_name']
                style.font.size = Pt(self.formatting_standards['font_size_headers'])
    
    def _add_contact_header(self, doc, candidate_data):
        """Add contact information header."""
        # Name (larger font)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(candidate_data['personal_info']['name'])
        name_run.font.name = self.formatting_standards['font_name']
        name_run.font.size = Pt(self.formatting_standards['font_size_name'])
        name_run.bold = True
        
        # Contact info
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info = f"{candidate_data['personal_info']['email']} | {candidate_data['personal_info'].get('phone', '')} | {candidate_data['personal_info'].get('location', '')}"
        contact_run = contact_para.add_run(contact_info)
        contact_run.font.name = self.formatting_standards['font_name']
        contact_run.font.size = Pt(self.formatting_standards['font_size_body'])
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_professional_summary(self, doc, summary):
        """Add professional summary section."""
        # Section header
        header_para = doc.add_paragraph()
        header_run = header_para.add_run("PROFESSIONAL SUMMARY")
        header_run.font.name = self.formatting_standards['font_name']
        header_run.font.size = Pt(self.formatting_standards['font_size_headers'])
        header_run.bold = True
        
        # Summary content
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(summary)
        summary_run.font.name = self.formatting_standards['font_name']
        summary_run.font.size = Pt(self.formatting_standards['font_size_body'])
        
        doc.add_paragraph()  # Spacing
    
    def _add_skills_section(self, doc, skills):
        """Add skills section."""
        # Section header
        header_para = doc.add_paragraph()
        header_run = header_para.add_run("TECHNICAL SKILLS")
        header_run.font.name = self.formatting_standards['font_name']
        header_run.font.size = Pt(self.formatting_standards['font_size_headers'])
        header_run.bold = True
        
        # Skills content
        for category, skill_list in skills.items():
            if skill_list:
                category_para = doc.add_paragraph()
                category_run = category_para.add_run(f"{category.replace('_', ' ').title()}: ")
                category_run.font.name = self.formatting_standards['font_name']
                category_run.font.size = Pt(self.formatting_standards['font_size_body'])
                category_run.bold = True
                
                skills_run = category_para.add_run(", ".join(skill_list))
                skills_run.font.name = self.formatting_standards['font_name']
                skills_run.font.size = Pt(self.formatting_standards['font_size_body'])
        
        doc.add_paragraph()  # Spacing
    
    def _add_experience_section(self, doc, experiences):
        """Add professional experience section."""
        # Section header
        header_para = doc.add_paragraph()
        header_run = header_para.add_run("PROFESSIONAL EXPERIENCE")
        header_run.font.name = self.formatting_standards['font_name']
        header_run.font.size = Pt(self.formatting_standards['font_size_headers'])
        header_run.bold = True
        
        for exp in experiences:
            # Job title and company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{exp['title']} | {exp['company']}")
            title_run.font.name = self.formatting_standards['font_name']
            title_run.font.size = Pt(self.formatting_standards['font_size_body'])
            title_run.bold = True
            
            # Dates and location
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"{exp['start_date']} - {exp['end_date']} | {exp.get('location', '')}")
            date_run.font.name = self.formatting_standards['font_name']
            date_run.font.size = Pt(self.formatting_standards['font_size_body'])
            date_run.italic = True
            
            # Bullet points
            for bullet in exp.get('bullet_points', []):
                bullet_para = doc.add_paragraph()
                bullet_para.style = 'List Bullet'
                bullet_run = bullet_para.add_run(bullet.lstrip('• '))
                bullet_run.font.name = self.formatting_standards['font_name']
                bullet_run.font.size = Pt(self.formatting_standards['font_size_body'])
        
        doc.add_paragraph()  # Spacing
    
    def _add_education_section(self, doc, education):
        """Add education section."""
        if not education:
            return
            
        # Section header
        header_para = doc.add_paragraph()
        header_run = header_para.add_run("EDUCATION")
        header_run.font.name = self.formatting_standards['font_name']
        header_run.font.size = Pt(self.formatting_standards['font_size_headers'])
        header_run.bold = True
        
        for edu in education:
            # Degree and institution
            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(f"{edu['degree']} | {edu['institution']}")
            edu_run.font.name = self.formatting_standards['font_name']
            edu_run.font.size = Pt(self.formatting_standards['font_size_body'])
            edu_run.bold = True
            
            # Graduation year and GPA
            if edu.get('graduation_year') or edu.get('gpa'):
                details_para = doc.add_paragraph()
                details_text = edu.get('graduation_year', '')
                if edu.get('gpa'):
                    details_text += f" | GPA: {edu['gpa']}"
                details_run = details_para.add_run(details_text)
                details_run.font.name = self.formatting_standards['font_name']
                details_run.font.size = Pt(self.formatting_standards['font_size_body'])
                details_run.italic = True