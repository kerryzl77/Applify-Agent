import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import datetime
import platform
import tempfile
import shutil

class OutputFormatter:
    def __init__(self):
        # Use tempfile for Heroku compatibility
        self.temp_dir = tempfile.mkdtemp()
        self.output_dir = os.path.join(self.temp_dir, 'output')
        
        # Ensure output directory exists
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            
        # Check if running in Docker or Heroku
        self.in_docker = os.path.exists('/.dockerenv')
        self.in_heroku = bool(os.environ.get('DYNO'))
        
        # Clean up old files periodically
        self._cleanup_old_files()
    
    def _cleanup_old_files(self):
        """Clean up files older than 1 hour."""
        try:
            current_time = datetime.datetime.now()
            for filename in os.listdir(self.output_dir):
                file_path = os.path.join(self.output_dir, filename)
                if os.path.isfile(file_path):
                    file_time = datetime.datetime.fromtimestamp(os.path.getmtime(file_path))
                    if (current_time - file_time).total_seconds() > 900:  # 15 minutes
                        os.remove(file_path)
        except Exception as e:
            print(f"Error cleaning up old files: {str(e)}")
    
    def format_text(self, content, content_type):
        """Return formatted plain text."""
        if content_type == 'linkedin_message':
            # Ensure LinkedIn message is under 200 characters
            if len(content) > 200:
                content = content[:197] + '...'
        
        elif content_type in ['connection_email', 'hiring_manager_email']:
            # Ensure emails are reasonable length (roughly 200 words)
            words = content.split()
            if len(words) > 220:  # Give a little buffer
                content = ' '.join(words[:200]) + '...'
        
        elif content_type == 'cover_letter':
            # Ensure cover letter is reasonable length (roughly 350 words)
            words = content.split()
            if len(words) > 370:  # Give a little buffer
                content = ' '.join(words[:350]) + '...'
        
        return content
    
    def create_docx(self, content, job_data, candidate_data, content_type):
        """Create a DOCX file with the formatted content."""
        try:
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            if content_type == 'cover_letter':
                self._format_cover_letter_docx(doc, content, job_data, candidate_data)
            else:
                # For emails, use a simpler format
                self._format_email_docx(doc, content, job_data, candidate_data, content_type)
            
            # Generate filename with user ID to prevent conflicts
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            company_name = job_data['company_name'].replace(' ', '_')
            filename = f"{content_type}_{company_name}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Save the document
            doc.save(filepath)
            
            # Return both the filename and the full path
            return {
                'filename': filename,
                'filepath': filepath
            }
            
        except Exception as e:
            print(f"Error creating document: {str(e)}")
            return None
    
    def convert_to_pdf(self, docx_info):
        """Convert DOCX to PDF and truncate to one page."""
        try:
            docx_path = docx_info['filepath']
            filename = docx_info['filename']
            
            # Generate PDF path
            pdf_filename = filename.replace('.docx', '.pdf')
            pdf_path = os.path.join(self.output_dir, pdf_filename)
            temp_pdf_path = pdf_path.replace('.pdf', '_temp.pdf')
            
            # Try different conversion methods in order of preference
            success = False
            
            # Method 1: Try using docx2pdf first
            try:
                convert(docx_path, temp_pdf_path)
                if os.path.exists(temp_pdf_path):
                    success = True
            except Exception as e:
                print(f"docx2pdf conversion failed: {str(e)}")
            
            # Method 2: Try using LibreOffice in Docker/Heroku
            if not success and (self.in_docker or self.in_heroku):
                try:
                    from subprocess import run, PIPE
                    cmd = ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                           self.output_dir, docx_path]
                    process = run(cmd, stdout=PIPE, stderr=PIPE)
                    if process.returncode == 0 and os.path.exists(pdf_path):
                        success = True
                    else:
                        print(f"LibreOffice conversion failed: {process.stderr.decode()}")
                except Exception as e:
                    print(f"LibreOffice conversion failed: {str(e)}")
            
            # Method 3: Try using PyPDF2 directly (fallback)
            if not success:
                try:
                    from PyPDF2 import PdfReader, PdfWriter
                    from docx2pdf import convert
                    
                    convert(docx_path, temp_pdf_path)
                    reader = PdfReader(temp_pdf_path)
                    writer = PdfWriter()
                    
                    if len(reader.pages) > 0:
                        writer.add_page(reader.pages[0])
                        with open(pdf_path, 'wb') as output_file:
                            writer.write(output_file)
                        success = True
                except Exception as e:
                    print(f"PyPDF2 conversion failed: {str(e)}")
            
            # Clean up temporary file
            if os.path.exists(temp_pdf_path):
                try:
                    os.remove(temp_pdf_path)
                except Exception as e:
                    print(f"Error removing temporary file: {str(e)}")
            
            if success and os.path.exists(pdf_path):
                return {
                    'filename': pdf_filename,
                    'filepath': pdf_path
                }
            
            return None
            
        except Exception as e:
            print(f"Error converting to PDF: {str(e)}")
            return None
    
    def _format_cover_letter_docx(self, doc, content, job_data, candidate_data):
        """Format a cover letter in DOCX."""
        try:
            # Add candidate info at the top
            candidate_info = doc.add_paragraph()
            candidate_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
            candidate_info.add_run(candidate_data['personal_info']['name']).bold = True
            candidate_info.add_run('\n' + candidate_data['personal_info']['email'])
            candidate_info.add_run('\n' + candidate_data['personal_info']['phone'])
            # if candidate_data['personal_info']['linkedin']:
            #     candidate_info.add_run('\n' + candidate_data['personal_info']['linkedin'])
            
            # Add date
            date = doc.add_paragraph()
            date.alignment = WD_ALIGN_PARAGRAPH.LEFT
            date.add_run(datetime.datetime.now().strftime("%B %d, %Y"))
            
            # Add space before content
            doc.add_paragraph()
            
            # Split content into paragraphs and add them
            paragraphs = content.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.add_run(paragraph.strip())
            
            # # Add signature space
            # doc.add_paragraph()  # Add space
            # signature = doc.add_paragraph()
            # signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        except Exception as e:
            print(f"Error formatting cover letter: {str(e)}")
            raise
    
    def _format_email_docx(self, doc, content, job_data, candidate_data, content_type):
        """Format an email in DOCX."""
        # Add subject line for emails
        if content_type == 'connection_email':
            subject = f"Connection Request - {candidate_data['personal_info']['name']}"
        elif content_type == 'hiring_manager_email':
            subject = f"Application for {job_data['job_title']} Position - {candidate_data['personal_info']['name']}"
        
        subject_line = doc.add_paragraph()
        subject_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subject_line.add_run('Subject: ').bold = True
        subject_line.add_run(subject)
        
        # Add content
        doc.add_paragraph()  # Add space
        for paragraph in content.split('\n'):
            if paragraph.strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(paragraph) 